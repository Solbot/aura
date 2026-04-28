# knowledge.py
# RAG engine: watches ~/knowledge/upload, extracts text, chunks, indexes via SQLite FTS5.

import os
import shutil
import time
import db

UPLOAD_DIR    = os.path.expanduser("~/knowledge/upload")
PROCESSED_DIR = os.path.expanduser("~/knowledge/processed")


def init_dirs():
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    os.makedirs(PROCESSED_DIR, exist_ok=True)


# --- Text extraction ---

def _extract_txt(filepath):
    with open(filepath, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf(filepath):
    try:
        import pypdf
        reader = pypdf.PdfReader(filepath)
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
        return "\n\n".join(parts)
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def _extract_docx(filepath):
    try:
        import docx
        doc = docx.Document(filepath)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Include table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n\n".join(parts)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def _extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(filepath)
    elif ext in (".docx",):
        return _extract_docx(filepath)
    elif ext in (".txt", ".md", ".rst", ".csv", ".json", ".html", ".htm", ".xml"):
        return _extract_txt(filepath)
    else:
        try:
            return _extract_txt(filepath)
        except Exception:
            return None


# --- Chunking ---

def _chunk_text(text, chunk_words=400, overlap_words=40):
    """Split text into word-based chunks with overlap."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = " ".join(words[i:i + chunk_words])
        chunks.append(chunk)
        if i + chunk_words >= len(words):
            break
        i += chunk_words - overlap_words
    return chunks


# --- Processing ---

def process_file(filepath):
    """Extract, chunk, and index a file. Returns (success, message)."""
    filename = os.path.basename(filepath)
    text = _extract_text(filepath)

    if not text or len(text.strip()) < 50:
        return False, f"Could not extract usable text from '{filename}'"

    chunks = _chunk_text(text)
    if not chunks:
        return False, f"No chunks produced from '{filename}'"

    doc_id = db.knowledge_doc_add(filename, filepath, len(chunks))
    for i, chunk in enumerate(chunks):
        db.knowledge_chunk_add(doc_id, i, chunk)

    # Move to processed, avoiding collisions
    dest = os.path.join(PROCESSED_DIR, filename)
    if os.path.exists(dest):
        base, ext = os.path.splitext(filename)
        dest = os.path.join(PROCESSED_DIR, f"{base}_{int(time.time())}{ext}")
    try:
        shutil.move(filepath, dest)
    except Exception:
        pass  # file stays in upload — not fatal

    return True, f"Imported '{filename}': {len(chunks)} chunks indexed"


def watch_once():
    """Check the upload directory and process any new files. Returns result messages."""
    results = []
    if not os.path.exists(UPLOAD_DIR):
        return results
    for filename in sorted(os.listdir(UPLOAD_DIR)):
        filepath = os.path.join(UPLOAD_DIR, filename)
        if not os.path.isfile(filepath):
            continue
        success, msg = process_file(filepath)
        results.append(msg)
    return results


def search(query, limit=5):
    """Search the knowledge base. Returns list of result dicts."""
    return db.knowledge_search(query, limit)
